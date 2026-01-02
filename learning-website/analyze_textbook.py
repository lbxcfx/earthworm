from docx import Document
import os
import re

book_dir = r"e:\earthworm\book"

# 存储所有单元的核心内容
all_units = {}

for filename in sorted(os.listdir(book_dir)):
    if not filename.endswith('.docx'):
        continue
    
    unit_match = re.search(r'Unit (\d+)', filename)
    if not unit_match:
        continue
    
    unit_num = int(unit_match.group(1))
    filepath = os.path.join(book_dir, filename)
    doc = Document(filepath)
    
    # 合并所有文本和表格内容
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += cell.text + " "
            full_text += "\n"
    
    # 提取单元标题
    title_match = re.search(r'5AU\d+\s+(.+?)(?:\n|$)', full_text)
    unit_title = title_match.group(1).strip() if title_match else "Unknown"
    
    print(f"\n{'='*70}")
    print(f"UNIT {unit_num}: {unit_title}")
    print('='*70)
    
    # 寻找所有英文单词和短语
    # 查找包含具体词汇的行
    vocab_patterns = [
        r'\b(fish|beef|noodle|milk|fruit|vegetable|candy|chocolate|cola|cake|chip)\b',
        r'\b(soup|chicken|apple|potato|meat|bread|egg|juice|water|bottle)\b',
        r'\b(paper|plastic|earth|tree|recycle|clean|waste|protect)\b',
        r'\b(friend|share|help|happy|kind|nice|together|play|listen|respect)\b',
        r'\b(try|learn|practice|goal|brave|keep|fail|success|dream|grow)\b',
        r'\b(future|robot|travel|fly|space|car|technology|smart|live|change)\b',
        r'\b(festival|celebrate|lantern|dragon|mooncake|firework|wish|luck)\b',
    ]
    
    found_words = set()
    for pattern in vocab_patterns:
        matches = re.findall(pattern, full_text.lower())
        found_words.update(matches)
    
    print("\n找到的词汇:")
    print(", ".join(sorted(found_words)))
    
    # 查找核心句型
    sentence_patterns = [
        r"(Don't[^.!?]+[.!?])",
        r"(Let's[^.!?]+[.!?])",
        r"(You should[^.!?]+[.!?])",
        r"(We should[^.!?]+[.!?])",
        r"(Will you[^.!?]+[.!?])",
        r"(I'm going to[^.!?]+[.!?])",
        r"(How do[^.!?]+[.!?])",
        r"(What do[^.!?]+[.!?])",
    ]
    
    found_sentences = set()
    for pattern in sentence_patterns:
        matches = re.findall(pattern, full_text, re.IGNORECASE)
        found_sentences.update(matches[:3])  # 只取前3个
    
    if found_sentences:
        print("\n核心句型示例:")
        for s in list(found_sentences)[:5]:
            print(f"  - {s.strip()}")
