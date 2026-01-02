from docx import Document
import os

book_dir = r"e:\earthworm\book"

for filename in sorted(os.listdir(book_dir)):
    if filename.endswith('.docx') and 'Unit 1' in filename:
        print(f"\n{'='*60}")
        print(f"Unit 1 - Eat Healthily")
        print('='*60)
        
        doc = Document(os.path.join(book_dir, filename))
        
        # 读取文档
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) > 3:
                print(text)
        
        # 也读取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    print(row_text)
