from docx import Document
import os
import re

book_dir = r"e:\earthworm\book"

units_data = {}

for filename in sorted(os.listdir(book_dir)):
    if filename.endswith('.docx'):
        unit_match = re.search(r'Unit (\d+)', filename)
        if not unit_match:
            continue
            
        unit_num = int(unit_match.group(1))
        filepath = os.path.join(book_dir, filename)
        doc = Document(filepath)
        
        # 合并所有文本
        all_text = ""
        for para in doc.paragraphs:
            all_text += para.text + " "
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        
        # 提取英文单词
        english_words = re.findall(r'\b[a-z]+\b', all_text.lower())
        # 过滤常见词
        common = {'the', 'a', 'an', 'is', 'are', 'be', 'to', 'of', 'and', 'in', 'on', 'at', 'for', 'with', 'it', 'this', 'that', 'can', 'will', 'do', 'does', 'don', 'should', 'let', 's', 'you', 'i', 'we', 'they', 'he', 'she', 'my', 'your', 'our', 'their', 'his', 'her', 'what', 'how', 'where', 'when', 'who', 'why', 'yes', 'no', 'not', 'if', 'or', 'but', 'so', 'up', 'about', 'more', 'some', 'all', 'any', 'each', 'other', 'new', 'old', 'good', 'bad', 'big', 'small', 'just', 'only', 'from', 'into', 'out', 'over', 'after', 'before', 'between', 'under', 'around', 'through', 'during', 'without', 'within', 'along', 'across', 'behind', 'beyond', 'by', 'down', 'off', 'than', 'then', 'there', 'here', 'now', 'very', 'much', 'many', 'too', 'also', 'as', 'was', 'were', 'been', 'being', 'have', 'has', 'had', 'having', 'make', 'makes', 'made', 'making', 'get', 'gets', 'got', 'getting', 'go', 'goes', 'went', 'going', 'come', 'comes', 'came', 'coming', 'see', 'sees', 'saw', 'seeing', 'take', 'takes', 'took', 'taking', 'think', 'thinks', 'thought', 'thinking', 'know', 'knows', 'knew', 'knowing', 'want', 'wants', 'wanted', 'wanting', 'use', 'uses', 'used', 'using', 'find', 'finds', 'found', 'finding', 'give', 'gives', 'gave', 'giving', 'tell', 'tells', 'told', 'telling', 'may', 'might', 'must', 'shall', 'would', 'could', 'one', 'two', 'three', 'four', 'five', 'first', 'second', 'third', 'last', 'next', 'way', 'day', 'time', 'year', 'week', 'month', 'read', 'look', 'say', 'said', 'talk', 'ask', 'answer', 'write', 'listen', 'watch', 'work', 'play', 'learn', 'teach', 'study', 'help', 'try', 'start', 'begin', 'end', 'finish', 'stop', 'keep', 'put', 'bring', 'show', 'turn', 'call', 'need', 'feel', 'become', 'leave', 'still', 'even', 'back', 'again', 'today', 'yesterday', 'tomorrow', 'morning', 'afternoon', 'evening', 'night', 'am', 'pm', 'well', 'really', 'able', 'like', 'likes', 'liked', 'liking'}
        
        # 统计词频
        word_freq = {}
        for word in english_words:
            if word not in common and len(word) > 2:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # 排序
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        
        print(f"\n{'='*60}")
        print(f"UNIT {unit_num}")
        print('='*60)
        
        # 找主题
        theme_match = re.search(r'"([A-Za-z\s]+)"', all_text)
        if theme_match:
            print(f"主题: {theme_match.group(1)}")
        
        # 打印高频词汇（可能是核心词汇）
        print("\n高频英文词汇 (可能的核心词汇):")
        vocab_words = []
        for word, count in sorted_words[:30]:
            if count >= 3:
                vocab_words.append(word)
        print(", ".join(vocab_words))
