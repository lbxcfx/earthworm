from docx import Document
import os

book_dir = r"e:\earthworm\book"

for filename in sorted(os.listdir(book_dir)):
    if filename.endswith('.docx'):
        print(f"\n\n{'#'*80}")
        print(f"# FILE: {filename}")
        print('#'*80)
        
        doc = Document(os.path.join(book_dir, filename))
        
        # 读取所有段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                print(text)
        
        print("\n--- TABLES ---")
        # 读取所有表格
        for ti, table in enumerate(doc.tables):
            print(f"\n[Table {ti+1}]")
            for row in table.rows:
                cells = [cell.text.strip()[:50] for cell in row.cells if cell.text.strip()]
                if cells:
                    print(" | ".join(cells))
