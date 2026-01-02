from docx import Document
import os
import re

book_dir = r"e:\earthworm\book"

for filename in sorted(os.listdir(book_dir)):
    if filename.endswith('.docx'):
        unit_num = re.search(r'Unit (\d+)', filename)
        if unit_num:
            print(f"\n{'#'*70}")
            print(f"# UNIT {unit_num.group(1)}")
            print('#'*70)
        
        doc = Document(os.path.join(book_dir, filename))
        
        full_text = ""
        for para in doc.paragraphs:
            full_text += para.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += cell.text + " | "
                full_text += "\n"
        
        # 查找【核心词汇】和【核心句型】
        vocab_match = re.findall(r'【核心词汇】[^\n【]*', full_text)
        sentence_match = re.findall(r'【核心句型】[^\n【]*', full_text)
        
        if vocab_match:
            print("\n=== 核心词汇 ===")
            for v in vocab_match[:3]:
                # 清理并打印
                clean = v.replace('【核心词汇】', '').strip()
                if clean:
                    print(clean)
        
        if sentence_match:
            print("\n=== 核心句型 ===")
            for s in sentence_match[:3]:
                clean = s.replace('【核心句型】', '').strip()
                if clean:
                    print(clean)
        
        # 查找单元主题
        theme_match = re.search(r'"([^"]+)"，本主题', full_text)
        if theme_match:
            print(f"\n主题: {theme_match.group(1)}")
