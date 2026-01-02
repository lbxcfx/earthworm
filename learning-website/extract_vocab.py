from docx import Document
import os
import re

book_dir = r"e:\earthworm\book"

for filename in sorted(os.listdir(book_dir)):
    if filename.endswith('.docx'):
        unit_num = re.search(r'Unit (\d+)', filename)
        if unit_num:
            print(f"\n{'='*70}")
            print(f"UNIT {unit_num.group(1)}: {filename}")
            print('='*70)
        
        doc = Document(os.path.join(book_dir, filename))
        
        full_text = ""
        for para in doc.paragraphs:
            full_text += para.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += cell.text + " "
                full_text += "\n"
        
        # 查找核心词汇部分
        if '核心词汇' in full_text or '词汇' in full_text:
            # 提取包含英文单词的行
            lines = full_text.split('\n')
            for line in lines:
                # 查找包含英文单词的行
                if re.search(r'[a-zA-Z]{3,}', line) and len(line) < 200:
                    text = line.strip()
                    if text and ('【' in text or '核心' in text or re.match(r'^[a-zA-Z]', text.strip()) or 'eat' in text.lower() or 'drink' in text.lower()):
                        print(text[:150])
